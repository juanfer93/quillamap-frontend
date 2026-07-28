from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\dev\quillamap-frontend")
OUT = ROOT / "docs" / "informe-avance-rutas-transporte-quillamap-2026-07-23.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr[index], header, bold=True)
        set_cell_shading(hdr[index], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


doc = Document()
configure_styles(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Informe de avance - Rutas de transporte publico en QuillaMap")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph("Fecha: 23 de julio de 2026 | Estado: trabajo local, no subido a GitHub")
subtitle.runs[0].italic = True

doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    "Se avanzo lentamente por bloques, como se definio para evitar errores en un trabajo pesado de datos. "
    "Ya quedaron inyectadas rutas oficiales de SOBUSA, Coolitoral y Cootrantico desde archivos KMZ publicados por el AMB. "
    "La cobertura validada subio de 34 a 66 rutas con shape valido y las rutas faltantes bajaron de 97 a 65."
)

add_table(
    doc,
    ["Metrica", "Base", "Despues de SOBUSA", "Despues de Coolitoral", "Despues de Cootrantico"],
    [
        ["Rutas esperadas", "131", "131", "131", "131"],
        ["Rutas con shape valido", "34", "43", "58", "66"],
        ["Rutas faltantes", "97", "88", "73", "65"],
        ["Shapes incompletos expuestos", "0", "0", "0", "0"],
        ["Paraderos OSM", "588", "588", "588", "588"],
    ],
    [1.65, 0.75, 1.25, 1.3, 1.45],
)

doc.add_heading("Lo que se hizo", level=1)
add_bullets(doc, [
    "Se creo y ejecuto un auditor de rutas contra /api/transit/routes/map para medir cobertura real.",
    "Se completo el catalogo maestro de rutas colectivas en el seed de transit.",
    "Se bloqueo la exposicion de shapes con menos de 3 puntos para evitar diagonales falsas.",
    "Se agrego un seed dedicado para SOBUSA con KMZ oficiales del AMB.",
    "Se agrego un seed dedicado para Coolitoral con KMZ oficiales del AMB.",
    "Se agrego y ejecuto un seed dedicado para Cootrantico con 8 KMZ oficiales del AMB.",
    "Se mantuvo el trabajo local, sin commit, push ni cambios de UI.",
])

doc.add_heading("SOBUSA completado", level=2)
doc.add_paragraph("Fuente oficial: https://www.ambq.gov.co/rutas-de-buses/SOBUSA/")
add_table(
    doc,
    ["Ruta", "Estado", "Puntos"],
    [
        ["B18-4175 A", "valid_shape", "1265"],
        ["B18-4175 B", "valid_shape", "1265"],
        ["C11-4168", "valid_shape", "1474"],
        ["C12-4169 A", "valid_shape", "860"],
        ["C12-4169 B", "valid_shape", "860"],
        ["C13-4143", "valid_shape", "875"],
        ["C14-4170", "valid_shape", "1290"],
        ["C16-4167 A", "valid_shape", "1590"],
        ["C16-4167 B", "valid_shape", "1590"],
    ],
    [2.2, 2.0, 1.0],
)

doc.add_heading("Coolitoral completado", level=2)
doc.add_paragraph("Fuente oficial: https://www.ambq.gov.co/ruta-de-buses/COOLITORAL/")
add_table(
    doc,
    ["Ruta", "Estado", "Puntos"],
    [
        ["A1-4106 A", "valid_shape", "541"],
        ["A1-4106 B", "valid_shape", "436"],
        ["A2-4107", "valid_shape", "601"],
        ["A3-4108", "valid_shape", "365"],
        ["A4-4109", "valid_shape", "288"],
        ["B1-4117", "valid_shape", "242"],
        ["B2A-4177", "valid_shape", "718"],
        ["B3-4119", "valid_shape", "256"],
        ["B17-4163", "valid_shape", "314"],
        ["C19-4178", "valid_shape", "497"],
        ["PT1", "valid_shape", "387"],
        ["PT2", "valid_shape", "964"],
        ["PT3", "valid_shape", "1007"],
        ["PT4", "valid_shape", "467"],
        ["PT5", "valid_shape", "467"],
    ],
    [2.2, 2.0, 1.0],
)

doc.add_heading("Cootrantico completado", level=2)
doc.add_paragraph("Fuente oficial: https://www.ambq.gov.co/ruta-de-buses/COOTRANTICO/")
add_table(
    doc,
    ["Ruta", "Estado", "Puntos"],
    [
        ["A18-4183", "valid_shape", "490"],
        ["B4-4120", "valid_shape", "300"],
        ["B5-4121", "valid_shape", "519"],
        ["B5-B-4190", "valid_shape", "645"],
        ["B6-4122", "valid_shape", "692"],
        ["B7-4123", "valid_shape", "634"],
        ["B20-4180", "valid_shape", "401"],
        ["B20-B-4191", "valid_shape", "505"],
    ],
    [2.2, 2.0, 1.0],
)

doc.add_heading("Comandos validados", level=1)
add_bullets(doc, [
    "Backend: npm.cmd run build",
    "Backend: npm.cmd test -- --runInBand src/features/transit/transit.service.spec.ts",
    "Backend: npm.cmd run seed:transit:sobusa",
    "Backend: npm.cmd run seed:transit:coolitoral",
    "Backend: npm.cmd run seed:transit:cootrantico",
    "Frontend auditoria: npm.cmd run audit:transit-routes -- --base-url http://127.0.0.1:3000/api --out-md docs\\transit-route-audit.latest.md --out-json docs\\transit-route-audit.latest.json",
    "Endpoint: http://127.0.0.1:3000/api/transit/routes/map respondio 200",
])

doc.add_heading("Lo que falta", level=1)
doc.add_paragraph(
    "Quedan 65 rutas esperadas sin shape valido en la auditoria. La prioridad sigue siendo avanzar por empresa, "
    "con fuente oficial cuando exista, convertir KMZ/KML a GeoJSON LineString y correr la auditoria despues de cada bloque."
)
add_table(
    doc,
    ["Empresa / sistema", "Rutas pendientes aproximadas", "Prioridad sugerida"],
    [
        ["Lolaya", "3", "Siguiente bloque recomendado por tener KMZ oficiales AMB"],
        ["Cootrasol", "3", "Bloque pequeno con KMZ oficiales"],
        ["Embusa", "1", "Bloque pequeno"],
        ["Transmecar", "4", "Tiene directorio historico AMB"],
        ["Resto de empresas", "54", "Avanzar lentamente por fuente y disponibilidad"],
    ],
    [2.2, 1.8, 2.5],
)

doc.add_heading("Siguiente paso recomendado", level=1)
add_numbered(doc, [
    "Cerrar la jornada con SOBUSA, Coolitoral y Cootrantico completados.",
    "Seguir con Lolaya como cuarto bloque en una proxima jornada.",
    "Crear seed dedicado o extender el importador oficial por operador.",
    "Ejecutar el seed, regenerar auditoria y confirmar que missingRoutes baje de 65.",
    "No tocar la UI hasta cerrar varios bloques de datos.",
    "No subir a GitHub hasta autorizacion explicita.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
